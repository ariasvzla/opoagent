import asyncio
import contextvars
import os
import re
import unicodedata
from pathlib import Path
from typing import Awaitable, Callable, Optional

import boto3
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)
RUNS_DIR = OUTPUT_DIR / "runs"
RUNS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR = OUTPUT_DIR / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)

_progress_callback: Optional[Callable[[str, str], Awaitable[None]]] = None
_workspace_dir: contextvars.ContextVar[Optional[Path]] = contextvars.ContextVar("workspace_dir", default=None)


def _safe_slug(value: str) -> str:
    raw = (value or "workspace").strip().lower()
    # Normalize unicode: é → e + combining accent, then strip accents
    raw = unicodedata.normalize("NFKD", raw)
    raw = "".join(c for c in raw if not unicodedata.combining(c))
    slug = re.sub(r"[^a-z0-9_-]+", "-", raw).strip("-")
    return slug or "workspace"


def get_output_dir() -> Path:
    active = _workspace_dir.get()
    if active is not None:
        active.mkdir(parents=True, exist_ok=True)
        return active
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def set_output_workspace(path: Optional[str]) -> str:
    if not path:
        _workspace_dir.set(None)
        return str(OUTPUT_DIR)

    workspace = Path(path)
    workspace.mkdir(parents=True, exist_ok=True)
    _workspace_dir.set(workspace)
    return str(workspace)


def create_run_workspace(run_id: str) -> str:
    run_dir = RUNS_DIR / _safe_slug(run_id)
    run_dir.mkdir(parents=True, exist_ok=True)
    return str(run_dir)


def create_tema_workspace(run_id: str, tema_label: str) -> str:
    run_dir = Path(create_run_workspace(run_id))
    tema_dir = run_dir / _safe_slug(tema_label)
    tema_dir.mkdir(parents=True, exist_ok=True)
    return str(tema_dir)


def set_progress_callback(callback: Optional[Callable[[str, str], Awaitable[None]]]) -> None:
    global _progress_callback
    _progress_callback = callback


def emit_progress(message: str, event_type: str = "stage") -> None:
    callback = _progress_callback
    if callback is None:
        return
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return
    loop.create_task(callback(event_type, message))


def extract_text_from_file(file_path: str) -> str:
    """Extract readable text from a local uploaded file."""
    path = Path(file_path)
    if not path.exists():
        return f"missing:{path}"

    def read_text_fallback() -> str:
        """Fallback for files that are not valid binaries despite their extension."""
        return path.read_text(encoding="utf-8", errors="replace")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".csv"}:
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        try:
            doc = Document(path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        except Exception:
            return read_text_fallback()

    if suffix == ".pdf":
        try:
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(page for page in pages if page.strip())
        except Exception:
            return read_text_fallback()

    return f"unsupported file type: {suffix or 'unknown'}"


def save_uploaded_file(filename: str, content: str) -> str:
    """Persist uploaded text content to disk so the agent can inspect it."""
    safe_name = Path(filename).name
    path = UPLOADS_DIR / safe_name
    path.write_text(content, encoding="utf-8")
    return str(path)


def read_uploaded_file(path: str) -> str:
    """Read an uploaded file from disk and return its text content."""
    return extract_text_from_file(path)


def save_section(filename: str, content: str) -> str:
    """Save a generated markdown section to disk with a safe filename."""
    safe_name = _safe_slug(Path(filename).stem) + Path(filename).suffix
    emit_progress(f"📝 Writing section {safe_name}...")
    path = get_output_dir() / safe_name
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    emit_progress(f"✅ Saved section {safe_name}.")
    return f"saved:{path}"


def read_section(filename: str) -> str:
    """Read a generated markdown section from disk."""
    path = get_output_dir() / filename
    if not path.exists():
        return f"missing:{path}"
    return path.read_text(encoding="utf-8")


def _add_inline_markdown(paragraph, text: str) -> None:
    """Render basic inline markdown (bold, italic, code) into a docx paragraph."""
    if not text:
        return

    # Split preserving markdown tokens: **bold**, *italic*, `code`
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue

        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            continue

        if part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue

        paragraph.add_run(part)


def _render_markdown_to_docx(document: Document, content: str) -> None:
    """Render markdown-ish content into a readable .docx structure."""
    in_code_block = False

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_paragraph = document.add_paragraph()
            run = code_paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped in {"---", "***", "___"}:
            document.add_paragraph("-" * 48)
            continue

        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[hashes:].strip()
            level = max(1, min(hashes, 4))
            heading = document.add_heading("", level=level)
            _add_inline_markdown(heading, heading_text)
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_markdown(paragraph, bullet_match.group(1))
            continue

        number_match = re.match(r"^\s*\d+[\.)]\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_markdown(paragraph, number_match.group(1))
            continue

        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{quote_text}")
            run.italic = True
            continue

        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, line)


def assemble_document(title: str, filenames: list[str], output_filename: str = "05-final-document.docx") -> str:
    """Assemble markdown files into a .docx document."""
    emit_progress("🧩 Assembling the final document from the generated sections...")

    document = Document()
    document.add_heading(title, level=1)

    output_dir = get_output_dir()
    added_sections = 0

    for name in filenames:
        raw_path = Path(name)
        path = raw_path if raw_path.is_absolute() else (output_dir / raw_path)
        if not path.exists():
            continue

        content = path.read_text(encoding="utf-8")
        if not content.strip():
            continue
        _render_markdown_to_docx(document, content)
        document.add_paragraph("")
        added_sections += 1

    if added_sections == 0:
        document.add_paragraph(
            "No se pudo incluir contenido en el documento. "
            "Revisa failures.md y los archivos tema-content.md/tema-tests.md del workspace."
        )

    final_path = output_dir / output_filename
    final_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(final_path)
    emit_progress(f"✅ Final document assembled at {final_path}.")
    return f"assembled:{final_path}"


def upload_document_to_s3(
    filename: str = "05-final-document.docx",
    bucket: Optional[str] = None,
    key: Optional[str] = None,
    run_id: Optional[str] = None,
) -> str:
    """Upload the assembled document to S3 with a meaningful path.

    When ``run_id`` is provided the key becomes:
        runs/{run_id}/{filename}

    Otherwise falls back to ``key`` or just the filename.
    """
    emit_progress(f"☁️ Preparing to upload {filename} to S3...")

    resolved_bucket = bucket or os.getenv("S3_BUCKET")
    if not resolved_bucket:
        emit_progress("⚠️ S3 upload skipped because S3_BUCKET is not configured.")
        return f"skipped:no-bucket:{filename}"

    path = get_output_dir() / filename
    if not path.exists():
        emit_progress(f"⚠️ S3 upload skipped because {path} does not exist.")
        return f"skipped:missing-file:{filename}"

    # Build a meaningful key: runs/{run_id}/{filename}
    if run_id:
        resolved_key = f"runs/{run_id}/{Path(filename).name}"
    else:
        resolved_key = key or filename
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "us-east-1"
    try:
        sts_client = boto3.client("sts", region_name=region)
        assumed_role_object = sts_client.assume_role(
            RoleArn=os.getenv("UPLOADER_ROLE_ARN"),
            RoleSessionName="agent-uploader"
        )
        credentials = assumed_role_object['Credentials']
        client = boto3.client(
            "s3",
            region_name=region,
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
        
        client.upload_file(str(path), resolved_bucket, resolved_key)
        publish_url = _generate_presigned_url(client, resolved_bucket, resolved_key)  # Generar URL pre-firmada para el archivo subido
        return f"uploaded:{publish_url}"
    except Exception as exc:
        emit_progress(f"⚠️ S3 upload failed: {exc}")
        return f"failed:{exc}"

def _generate_presigned_url(
    client,
    resolved_bucket: str,
    resolved_key: str,
    expiration: int = 86400,
) -> str:
    """
    Generate a pre-signed URL for downloading an S3 object.

    :param resolved_bucket: S3 bucket name
    :param resolved_key: S3 object key (path/filename)
    :param expiration: URL validity in seconds (default: 24 hours)
    :return: Pre-signed URL
    """
    return client.generate_presigned_url(
        ClientMethod="get_object",
        Params={
            "Bucket": resolved_bucket,
            "Key": resolved_key,
        },
        ExpiresIn=expiration,
    )